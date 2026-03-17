import os
import uuid
from typing import Sequence

import fitz
import numpy as np
from paddleocr import PaddleOCR
from tqdm import tqdm

from core.custom_types import DocumentMetadata
from llama_index.core.schema import Document as LlamaDocument
from docx import Document
import re

from core.settings import settings


def load_txt(path,metadata: DocumentMetadata)->Sequence[LlamaDocument]:
    """获取普通文本Document"""
    with open(path, "r", encoding="utf-8") as f:
        text = f.read()
    metadata.source = metadata.file_type
    metadata.section_title = ".".join(metadata.file_name.split(".")[:-1])
    return [
        LlamaDocument(
            text=text,
            metadata=metadata.dict()
        )
    ]

def load_docx(file_path: str, metadata: DocumentMetadata)->Sequence[LlamaDocument]:
    doc = Document(file_path)
    sections = []
    current_section = ""
    current_title = None
    metadata.source = metadata.file_type
    for para in doc.paragraphs:
        text = para.text.strip()

        if not text:
            continue

        # 如果是标题
        if para.style.name.startswith("Heading"):

            if current_section:
                sections.append((current_title,current_section))

            current_title = text
            current_section = f"# {text}\n"
        else:
            current_section += text + "\n"

    if current_section:
        sections.append((current_title,current_section))

    # 解析表格
    for table in doc.tables:

        table_text = []

        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells]

            table_text.append(" | ".join(row_data))

        sections.append((current_title,"\n".join(table_text)))

    documents = []

    for title,section in sections:
        metadata.section_title = title
        documents.append(
            LlamaDocument(
                text=section,
                metadata=metadata.dict()
            )
        )

    return documents



def extract_md_title(text: str) -> str | None:
    match = re.search(r'^(#{1,6})\s+(.*)', text, re.MULTILINE)
    if match:
        return match.group(2).strip()
    return None

def load_markdown(path: str, metadata: DocumentMetadata) -> Sequence[LlamaDocument]:
    """获取markdown Document"""
    with open(path, "r", encoding="utf-8") as f:
        text = f.read()

    sections = []
    # 按 markdown 标题切分
    parts = re.split(r'(?=\n#{1,6} )', text)

    metadata.source = metadata.file_type
    for part in parts:
        part = part.strip()
        if not part:
            continue
        metadata.section_title = extract_md_title(part)
        sections.append(
            LlamaDocument(
                text=part,
                metadata=metadata.dict()
            )
        )

    return sections



ocr = PaddleOCR(
    use_angle_cls=True,
    lang=settings.orc_lang,
    ocr_version='PP-OCRv4'
)

def extract_pdf_title(text: str) -> bool:
    """
    简单标题识别
    """
    if len(text) < 40 and re.match(r"^[0-9一二三四五六七八九十\.、 ]+", text):
        return True

    if len(text) < 20 and text.isupper():
        return True

    return False

def ocr_page(page):
    pix = page.get_pixmap(matrix=fitz.Matrix(3, 3))

    img_path = f"temp_{uuid.uuid4().hex}.png"
    pix.save(img_path)

    result = ocr.ocr(img_path)
    print(result)
    texts = []

    # 删除临时文件（很重要）
    try:
        os.remove(img_path)
    except:
        pass
    return "\n".join(texts)



def load_pdf(path: str, metadata: DocumentMetadata) -> Sequence[LlamaDocument]:
    """获取pdf Document"""
    pdf = fitz.open(path)
    documents = []

    current_section = ""
    current_title = ""
    source = metadata.file_type
    for page_num, page in tqdm(enumerate(pdf),desc="加载pdf"):

        blocks  = page.get_text("blocks")
        source =  metadata.file_type
        if len(blocks)<=1 or  sum(len(b[4]) for b in blocks) < 50:
            print(f"第{page_num+1}页扫描")
            #进行ocr扫描图文
            text = ocr_page(page)
            text = text.replace("\n", " ")
            blocks = [(0, 0, 0, 0, text, 0, 0)]
            source = "ocr"
        for block in blocks:
            text = block[4].strip()
            if not text:
                continue
            text = text.replace("\n", " ")

            # 标题、页码、页脚
            if len(text) < 5:
                continue

            if extract_pdf_title(text):
                if current_section:
                    documents.append(
                        LlamaDocument(
                            text=current_section,
                            metadata={
                                **metadata.dict(),
                                "page": page_num + 1,
                                "section_title": current_title,
                                "source":source
                            }
                        )
                    )
                current_title = text
                current_section = f"# {text}\n"
            else:
                current_section += text + "\n"

    if current_section:
        metadata.section_title =  current_title
        metadata.source =  metadata.file_type
        documents.append(
            LlamaDocument(
                text=current_section,
                metadata=metadata.dict()
            )
        )
    return documents




def load_file(path:str, metadata: DocumentMetadata)->Sequence[LlamaDocument]:
    if path.endswith(".txt"):
        return load_txt(path,metadata)
    elif path.endswith(".docx") or path.endswith(".doc"):
        return load_docx(path,metadata)
    elif path.endswith(".md"):
        return load_markdown(path,metadata)
    elif path.endswith(".pdf"):
        return load_pdf(path,metadata)
    else:
        return []


